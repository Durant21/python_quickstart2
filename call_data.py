import random
import docx

import mammoth

# from DocTrace_v3.BLL.Groups import BLL_Groups
from BLL.Documents import BLL_Documents
from Domain.Documents import Document
from data.db_factory import DbSessionFactory

from BLL.Sections import BLL_Sections
from BLL.Groups import BLL_Groups
from BLL.Doc_Parent import BLL_Doc_Parent
from BLL.DocGroupSections import BLL_DocGroupSections

# def test_call_groups():
#    r = BLL_Groups.get_groups()
#    return r
#
# def test_call_documents():
#    r = BLL_Documents.get_documents()
#    return r
#
# def test_doc():
#    # doc = Document('','')
#    # doc.__init__('','')
#    doc = Document()
#    doc.add_name('a doc')
#    print(doc.doc_name)
#    return doc
#

def test_get_all_documents():
    docs = BLL_Documents.get_documents()
    print('{} docs found'.format(len(docs)))

def test_create_doc():
    doc_data = {"doc_name": "new doc 2:22"}
    r = BLL_Documents.create_document(doc_data=doc_data)
    print(r)
#
def test_update_doc(doc_name):
    doc_data = {"doc_id": "1001","doc_name": "changed name2"}
    r = BLL_Documents.update_document(doc_id='1001',doc_data=doc_data)
    print(r)

def test_get_doc_by_name():
    doc_name = 'abc'
    doc = BLL_Documents.single_document_by_name(doc_name)
    if not doc:
        print('could not find {}'.format(doc_name))


def test_create_section():
    # section_data = request.json_body
    section_data = {"sec_date_in":"1/1/2019", "sec_text": "new section 1"}
    sec_id = BLL_Sections.create_section(section_data)
    print(sec_id)

    return sec_id

def test_get_all_sections():
    sections = BLL_Sections.get_sections()
    print('found {} sections'.format(len(sections)))
    return None

def test_get_sections_by_doc(doc_id):
    sections = BLL_Sections.get_sections_by_doc(doc_id)
    print('found {} sections'.format(len(sections)))
    return None

def test_edit_section():
    section_data = {
        "sec_id": "8c064d9a-f689-4ab1-bbed-1652623aa1c7",
        "sec_text": "new section 1a",
        "sec_date_in": "2019-01-01"
    }
    r = BLL_Sections.update_section("8c064d9a-f689-4ab1-bbed-1652623aa1c7", section_data)
    print(r)

def test_create_group(inc_num):
    # create a document
    doc_data = {"doc_name": "new doc 2:22_" + inc_num}
    r = BLL_Documents.create_document(doc_data=doc_data)
    doc_id = r["msg"]
    print("doc_id = {}".format(doc_id))
    # create a section
    section_data = {"sec_date_in":"1/1/2019", "sec_text": "new section 1"}
    r = BLL_Sections.create_section(section_data)
    sec_id = r["msg"]
    print("sec_id = {}".format(sec_id))

    # create a group
    group_data = {"doc_id": doc_id,
                  "sec_id": sec_id}
    r = BLL_Groups.create_group(group_data)
    group_id = r["msg"]
    print("created group" + group_id)

def test_create_doc_group_relationships():
    # get a random existing document
    docs = BLL_Documents.get_documents()
    y = random.randint(1,len(docs))
    aDoc = docs[y]
    aDoc_dict = aDoc.to_dict()
    doc_id = aDoc_dict['doc_id']

    y = random.randint(1,len(docs))
    aDoc = docs[y]
    aDoc_dict = aDoc.to_dict()
    append_doc_id = aDoc_dict['doc_id']


    # create a section
    # get random existing section
    sections = BLL_Sections.get_sections()
    u = random.randint(1,len(sections))
    aSection = sections[u]
    aSection_dict = aSection.to_dict()
    sec_id = aSection_dict['sec_id']

    # # create a group
    # group_data = {"doc_id": doc_id,
    #               "sec_id": sec_id}
    # r = BLL_Groups.create_group(group_data)
    # group_id = r["msg"]
    # print("created group" + group_id)

    j_body = {"doc_id": doc_id,
                  "append_doc_id": append_doc_id}
    r = BLL_DocGroupSections.attach_sections(j_body)

def test_create_doc_parent():
    doc_data = {"doc_id": "1001",
                "parent_id": "1002",
                "relationship": "copy"}

    r = BLL_Doc_Parent.create_doc_parent(doc_data)
    print(r)

def test_create_multiple_doc_parents():

    # get a random existing document
    docs = BLL_Documents.get_documents()
    y = random.randint(1,len(docs))
    aDoc = docs[y]
    aDoc_dict = aDoc.to_dict()
    doc_id = aDoc_dict['doc_id']

    y = random.randint(1,len(docs))
    aDoc = docs[y]
    aDoc_dict = aDoc.to_dict()
    parent_id = aDoc_dict['doc_id']

    doc_data = {"doc_id": doc_id,
                "parent_id": parent_id,
                "relationship": "copy"}

    r = BLL_Doc_Parent.create_doc_parent(doc_data)
    print(r)


def test_get_doc_parent_report():
    doc_id =  '8dd7e778-b79f-4055-8379-f614d084b4e4'
    doc_parent_relationship = BLL_Doc_Parent.get_doc_parent_by_doc_id(doc_id=doc_id)
    print(doc_parent_relationship)

def test_find_parent_child():
    docs_all = []

    aDict = {'relationship': 'direct',
             'doc_name': '',
             'parent_doc_name': '' ,
             'parent_doc_id': '0',
             'doc_id': '1'
    }
    docs_all.append(aDict)

    aDict = {'relationship': 'direct',
             'doc_name': '',
             'parent_doc_name': '' ,
             'parent_doc_id': '0',
             'doc_id': '2'
    }
    docs_all.append(aDict)

    aDict = {'relationship': 'direct',
             'doc_name': '',
             'parent_doc_name': '' ,
             'parent_doc_id': '1',
             'doc_id': '3'
    }
    docs_all.append(aDict)

    aDict = {'relationship': 'direct',
             'doc_name': '',
             'parent_doc_name': '' ,
             'parent_doc_id': '1',
             'doc_id': '4'
    }
    docs_all.append(aDict)

    aDict = {'relationship': 'direct',
             'doc_name': '',
             'parent_doc_name': '' ,
             'parent_doc_id': '5',
             'doc_id': '4'
    }
    docs_all.append(aDict)

    aDict = {'relationship': 'direct',
             'doc_name': '',
             'parent_doc_name': '' ,
             'parent_doc_id': '0',
             'doc_id': '5'
    }
    docs_all.append(aDict)

    aDict = {'relationship': 'direct',
             'doc_name': '',
             'parent_doc_name': '' ,
             'parent_doc_id': '4',
             'doc_id': '6'
    }
    docs_all.append(aDict)
    t = []
    doc_list = ['1']
    final_lst = []
    r = BLL_Doc_Parent.find_parent(doc_list,docs_all,final_lst)
    # print(parent_lst)
    final_lst3 = []
    t = BLL_Doc_Parent.find_child(doc_list,docs_all,final_lst3)
    print(t)


def test_create_multiple_docs():
    range_start = 1
    range_stop = 3

    for i in range(range_start,range_stop):
        doc_data = {"doc_name": "doc" +str(i)}
        r1 = BLL_Documents.create_document(doc_data=doc_data)
        for t in range(range_start,range_stop):
            doc_data = {"doc_name": "doc" + str((i * 10) + t)}
            r2 = BLL_Documents.create_document(doc_data=doc_data)
            #
            # doc_rel_data1 = {"doc_id": r2['msg'],
            #             "parent_id": r1['msg'],
            #             "relationship": "copy"}
            #
            # s = BLL_Doc_Parent.create_doc_parent(doc_rel_data1)

            j_body = {"doc_id": r2['msg'],
                      "append_doc_id": r1['msg']}
            r = BLL_DocGroupSections.attach_sections(j_body)

            for g in range(range_start,range_stop):
                doc_data = {"doc_name": "doc" + str((i * 100) + g)}
                r3 = BLL_Documents.create_document(doc_data=doc_data)

                j_body = {"doc_id": r3['msg'],
                          "append_doc_id": r2['msg']}
                r = BLL_DocGroupSections.attach_sections(j_body)

                for h in range(range_start, range_stop):
                    doc_data = {"doc_name": "doc" + str((i * 1000) + h)}
                    r4 = BLL_Documents.create_document(doc_data=doc_data)

                    j_body = {"doc_id": r4['msg'],
                              "append_doc_id": r3['msg']}
                    r = BLL_DocGroupSections.attach_sections(j_body)

                    for q in range(range_start, range_stop):
                        doc_data = {"doc_name": "doc" + str((i * 10000) + q)}
                        r5 = BLL_Documents.create_document(doc_data=doc_data)

                        j_body = {"doc_id": r5['msg'],
                                  "append_doc_id": r4['msg']}
                        r = BLL_DocGroupSections.attach_sections(j_body)

                        for k in range(range_start, range_stop):
                            doc_data = {"doc_name": "doc" + str((i * 10000) + k)}
                            r6 = BLL_Documents.create_document(doc_data=doc_data)

                            j_body = {"doc_id": r6['msg'],
                                      "append_doc_id": r5['msg']}
                            r = BLL_DocGroupSections.attach_sections(j_body)





def test_getText(filename):
    doc = docx.Document(filename)

    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def test_getHTML(filename):

    with open(filename, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        html = result.value  # The generated HTML
        messages = result.messages  # Any messages, such as warnings during conversion
        print(html)
        print(messages)
def main():
    # DbSessionFactory.global_init('myDocuments.sqlite')
    DbSessionFactory.global_init('pypi.sqlite')
    # (test_call_groups())
    # test_call_documents()
    # test_get_all_documents()
    # (test_doc())
    # test_create_doc('abc')
    # test_update_doc('abc')

    # TEST DOCUMENTS
    # test_create_doc()

    # TEST SECTIONS
    # test_create_section()
    # test_get_all_sections()
    # test_get_sections_by_doc('d60a7e19-6c2c-4d90-92c9-cb8fb6c189a0')
    test_edit_section()

    # TEST GROUPS
    # for t in range(1,250):
    #     test_create_group(str(t))
    # test_get_doc_group()

    # TEST DOC_PARENT
    # test_create_doc_parent()
    # test_get_doc_parent_report()
    # test_find_parent_child()
    #
    # for _ in range(1,20):
    #     test_create_doc_group_relationships()
    #
    # for _ in range(1,6):
    #     test_create_multiple_doc_parents()

    #
    # for _ in range (1,50000):
    #     r  = 5
    #     t = 4
    #     u = r/t
    #     print(u)


    # test_create_multiple_docs()

    # test_get_doc_by_name()

    # test conversion of docx to an array
    # test_getText("fern.docx")

    # test conversion of docx to HTML
    # test_getHTML("testdoc1.docx")

    print('done')

if __name__ == "__main__":
    main()


